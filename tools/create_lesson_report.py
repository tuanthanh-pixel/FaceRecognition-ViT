from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = PROJECT_ROOT / "outputs" / "report_assets"
OUTPUT_PATH = PROJECT_ROOT / "Tong_hop_Custom_Vision_Transformer_SIC.docx"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_REGULAR = Path("C:/Windows/Fonts/arial.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/arialbd.ttf")


def get_font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_wrapped_center(draw, box, text, fill="#12355B", size=26, bold=False):
    x1, y1, x2, y2 = box
    font = get_font(size, bold)
    max_chars = max(10, int((x2 - x1) / (size * 0.55)))
    wrapped = "\n".join(textwrap.wrap(text, width=max_chars))
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=7, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        wrapped,
        font=font,
        fill=fill,
        spacing=7,
        align="center",
    )


def draw_box(draw, coords, text, fill, size=25):
    draw.rounded_rectangle(coords, radius=18, fill=fill, outline="#355070", width=3)
    draw_wrapped_center(draw, coords, text, size=size, bold=True)


def draw_arrow(draw, start, end, color="#355070", width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    for offset in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + offset),
            end[1] + length * math.sin(angle + offset),
        )
        draw.line([end, point], fill=color, width=width)


def save_pipeline_diagram():
    path = ASSET_DIR / "pipeline_facevit.png"
    image = Image.new("RGB", (1800, 520), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 28), "PIPELINE NHẬN DIỆN KHUÔN MẶT HIỆN TẠI", font=get_font(34, True), fill="#12355B")
    items = [
        ("Ảnh RGB\nVGGFace2", "#D8F3DC"),
        ("Resize 224x224\nAugmentation", "#CAF0F8"),
        ("FaceViT\nViT-Tiny/16 style", "#FFE5B4"),
        ("Embedding 128 chiều\nL2 norm = 1", "#FFD6E0"),
        ("Euclidean Distance\nVerification / Identification", "#E2D1F9"),
    ]
    width, gap, y1, y2 = 295, 60, 145, 385
    x = 45
    for index, (label, color) in enumerate(items):
        draw_box(draw, (x, y1, x + width, y2), label, color)
        if index < len(items) - 1:
            draw_arrow(draw, (x + width, 265), (x + width + gap - 10, 265))
        x += width + gap
    image.save(path)
    return path


def save_model_diagram():
    path = ASSET_DIR / "facevit_shapes.png"
    image = Image.new("RGB", (1600, 1120), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 28), "FACE VISION TRANSFORMER - LUỒNG SHAPE", font=get_font(34, True), fill="#12355B")
    items = [
        ("Input\n[B, 3, 224, 224]", "#D8F3DC"),
        ("Patch Embedding\nConv2d 16x16, stride 16", "#CAF0F8"),
        ("196 patch tokens\n[B, 196, 192]", "#ADE8F4"),
        ("Thêm CLS + position\n[B, 197, 192]", "#FFE5B4"),
        ("12 Transformer Blocks\nLayerScale + DropPath", "#FFD6E0"),
        ("Final LayerNorm\nCLS [B, 192]", "#E2D1F9"),
        ("Linear 192 -> 128\nL2 normalize", "#CAFFBF"),
    ]
    y = 110
    for index, (label, color) in enumerate(items):
        draw_box(draw, (345, y, 1255, y + 105), label, color, size=25)
        if index < len(items) - 1:
            draw_arrow(draw, (800, y + 105), (800, y + 145))
        y += 150
    image.save(path)
    return path


def save_training_diagram():
    path = ASSET_DIR / "training_losses.png"
    image = Image.new("RGB", (1700, 680), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 28), "HAI NHÁNH HUẤN LUYỆN TRONG REPO", font=get_font(34, True), fill="#12355B")
    draw_box(draw, (90, 160, 525, 390), "train.py\nSemi-Hard Triplet Loss\nmargin = 0.2", "#CAF0F8", size=25)
    draw_box(draw, (640, 160, 1065, 390), "train_infonce.py\nInfoNCE / SupCon\ntemperature = 0.07", "#FFE5B4", size=25)
    draw_box(draw, (1190, 160, 1605, 390), "Checkpoint tốt nhất\nhistory.json\ntraining_curves.png", "#D8F3DC", size=25)
    draw_arrow(draw, (525, 275), (630, 275))
    draw_arrow(draw, (1065, 275), (1180, 275))
    draw.text((95, 485), "Cả hai nhánh dùng cùng FaceViT, cùng PKBatchSampler và cùng split identity.", font=get_font(28, True), fill="#355070")
    draw.text((95, 535), "Sau huấn luyện, test.py đánh giá verification và identification trên identity chưa thấy trong train.", font=get_font(26), fill="#355070")
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)


def add_table(document, headers, rows, widths_cm):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.bold = True
            run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
    document.add_paragraph()
    return table


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.right_indent = Cm(0.45)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def style_document(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 24, "12355B"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build_document():
    pipeline = save_pipeline_diagram()
    model_shapes = save_model_diagram()
    training = save_training_diagram()

    doc = Document()
    style_document(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BÁO CÁO MÔ HÌNH FACEVIT\nNHẬN DIỆN KHUÔN MẶT")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Samsung Innovation Campus - cập nhật theo mã nguồn hiện tại")
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(53, 80, 112)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Ngày cập nhật: 09/08/2026\n").bold = True
    meta.add_run("Model: FaceVisionTransformer, train từ đầu, không dùng pretrained weights\n")
    meta.add_run("Đầu ra chính: embedding khuôn mặt 128 chiều, chuẩn hóa L2")

    doc.add_picture(str(pipeline), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("1. Tóm tắt hiện trạng", level=1)
    doc.add_paragraph(
        "Dự án hiện tại không còn là bài toán phân loại closed-set bằng classifier Softmax trên Pins Face Recognition. "
        "Mã nguồn đang tập trung vào học embedding khuôn mặt bằng một Vision Transformer tự cài đặt. "
        "Embedding sau cùng được chuẩn hóa L2 và được so sánh bằng Euclidean Distance để phục vụ verification hoặc identification."
    )
    add_table(
        doc,
        ["Hạng mục", "Giá trị hiện tại"],
        [
            ("Dataset mặc định", "dataset/vggface2"),
            ("Train/Val/Test", "Tách theo identity: train từ thư mục train; val/test chia từ thư mục val"),
            ("Input", "Ảnh RGB resize về 224 x 224"),
            ("Backbone", "FaceViT, patch 16 x 16, embed_dim=192, depth=12, num_heads=3"),
            ("Output", "Embedding 128 chiều, L2 norm = 1"),
            ("Số tham số", "5,553,728 tham số theo src/model.py"),
            ("Loss hỗ trợ", "Semi-Hard Triplet Loss và InfoNCE/Supervised Contrastive"),
        ],
        [4.2, 11.8],
    )

    doc.add_heading("2. Kiến trúc FaceViT", level=1)
    doc.add_picture(str(model_shapes), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Ảnh 224 x 224 được chia thành 14 x 14 = 196 patch. Mỗi patch được chiếu sang vector 192 chiều bằng Conv2d với kernel_size=16 và stride=16. "
        "Model thêm một CLS token và position embedding, sau đó đưa chuỗi 197 token qua 12 Transformer Encoder Block."
    )
    add_code(
        doc,
        "[B, 3, 224, 224]\n"
        "-> Conv2d(3, 192, kernel=16, stride=16)\n"
        "-> [B, 196, 192]\n"
        "-> thêm CLS + position embedding: [B, 197, 192]\n"
        "-> 12 Transformer blocks\n"
        "-> CLS [B, 192] -> Linear(192, 128) -> L2 normalize",
    )
    doc.add_paragraph(
        "Mỗi Transformer block dùng kiến trúc Pre-Norm: LayerNorm trước Multi-Head Self-Attention và trước FeedForward. "
        "So với báo cáo cũ, block hiện tại có thêm LayerScale với init 1e-5 và DropPath tăng dần từ 0 đến 0.1 theo depth, giúp regularize khi mô hình sâu hơn."
    )

    doc.add_heading("3. Dữ liệu và DataLoader", level=1)
    doc.add_paragraph(
        "Hàm read_data yêu cầu VGGFace2 có hai thư mục con train và val. Identity trong hai thư mục này không được trùng nhau. "
        "Các identity trong thư mục val được shuffle bằng seed=42 rồi chia thành validation và test theo validation_identity_ratio, mặc định là 0.5."
    )
    add_table(
        doc,
        ["Thành phần", "Mô tả"],
        [
            ("FaceImageDataset", "Đọc ảnh bằng PIL, convert RGB, áp dụng transform và trả về image/label."),
            ("Transform train", "Resize 224, RandomHorizontalFlip, RandomRotation 8 độ, ColorJitter, ToTensor, Normalize mean/std 0.5."),
            ("Transform eval", "Resize 224, ToTensor, Normalize mean/std 0.5."),
            ("PKBatchSampler", "Lấy P identity và K ảnh mỗi identity để trong batch luôn có positive và negative."),
            ("Mặc định P x K", "identities_per_batch=8, images_per_identity=4, batch mining thực tế là 32 ảnh."),
            ("batch_size", "Mặc định 8, dùng cho val_images/test_images khi trích xuất embedding tuần tự."),
        ],
        [4.2, 11.8],
    )

    doc.add_heading("4. Huấn luyện", level=1)
    doc.add_picture(str(training), width=Inches(6.55))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Repo hiện có hai entry point huấn luyện. train.py dùng TripletMarginLoss với semi-hard mining online trong từng batch. "
        "train_infonce.py dùng InfoNCELoss/Supervised Contrastive trên embedding và nhãn trong batch. Cả hai đều dùng AdamW, Early Stopping và lưu checkpoint tốt nhất theo validation loss."
    )
    add_table(
        doc,
        ["Lệnh", "Ý nghĩa"],
        [
            ("python train.py", "Huấn luyện FaceViT bằng Semi-Hard Triplet Loss."),
            ("python train_infonce.py --loss_type supcon", "Huấn luyện bằng InfoNCE/SupCon. Tham số loss_type hiện chủ yếu là cờ cấu hình; script riêng quyết định loss."),
            ("python train.py --epochs 100 --experiment_name sic_facevit_vggface2_semi_hard", "Chạy experiment mặc định và lưu checkpoint theo tên experiment."),
            ("python test.py --experiment_name sic_facevit_vggface2_semi_hard", "Đánh giá checkpoint tốt nhất trên test identities."),
        ],
        [6.2, 9.8],
    )
    doc.add_paragraph(
        "Trong Triplet branch, positive được chọn là ảnh cùng identity xa anchor nhất. Negative ưu tiên semi-hard theo điều kiện PosDist < NegDist < PosDist + margin. "
        "Nếu không có semi-hard negative, code dùng negative xa nhất làm fallback để tránh ép model học mẫu quá cực đoan ngay từ đầu."
    )

    doc.add_heading("5. Cấu hình chính", level=1)
    add_table(
        doc,
        ["Tham số", "Mặc định", "Ghi chú"],
        [
            ("image_size", "224", "Kích thước ảnh đầu vào."),
            ("patch_size", "16", "Tạo 196 patch cho ảnh 224 x 224."),
            ("embed_dim", "192", "Số chiều token Transformer."),
            ("depth", "12", "Số Transformer Encoder Block."),
            ("num_heads", "3", "Mỗi head có 64 chiều."),
            ("mlp_ratio", "4.0", "FeedForward 192 -> 768 -> 192."),
            ("dropout", "0.1", "Dropout trong embedding, attention và MLP."),
            ("face_embedding_dim", "128", "Kích thước embedding đầu ra."),
            ("identities_per_batch", "8", "P trong PK sampler."),
            ("images_per_identity", "4", "K trong PK sampler."),
            ("triplet_margin", "0.2", "Margin cho Triplet Loss."),
            ("temperature", "0.07", "Nhiệt độ cho InfoNCE."),
            ("lr", "0.0003", "Learning rate của AdamW."),
            ("weight_decay", "0.0001", "Regularization."),
            ("early_stop", "10", "Patience của Early Stopping."),
        ],
        [4.0, 3.0, 9.0],
    )

    doc.add_heading("6. Đánh giá mô hình", level=1)
    doc.add_paragraph(
        "Sau khi train xong, test.py nạp checkpoint tốt nhất, trích xuất embedding của test identities, sau đó đánh giá hai nhóm bài toán."
    )
    add_table(
        doc,
        ["Bài toán", "Cách đánh giá"],
        [
            ("Verification", "Tạo 10,000 cặp ảnh cân bằng same/different, tính distance và báo cáo ROC-AUC, EER, accuracy tại EER, TAR@FAR."),
            ("Identification", "Chọn 5 ảnh mỗi identity làm gallery, phần còn lại làm probe; tính Recall@1, Recall@5 và mAP."),
        ],
        [4.2, 11.8],
    )
    add_bullet(doc, "ROC-AUC càng gần 1 càng tốt; 0.5 tương đương ngẫu nhiên.")
    add_bullet(doc, "EER càng thấp càng tốt; EER distance threshold có thể dùng làm ngưỡng ban đầu cho demo.")
    add_bullet(doc, "TAR@FAR=1% và 0.1% quan trọng khi dùng nhận diện cho điểm danh hoặc kiểm soát ra vào.")
    add_bullet(doc, "Recall@1 đo việc identity đúng đứng đầu danh sách; mAP đánh giá toàn bộ thứ tự gallery.")

    doc.add_heading("7. Artifact sinh ra", level=1)
    add_table(
        doc,
        ["Đường dẫn", "Vai trò"],
        [
            ("checkpoints/<experiment>_best.pth", "Checkpoint tốt nhất theo validation loss, chứa model_state_dict, optimizer_state_dict, config và split identity."),
            ("outputs/<experiment>/history.json", "Lịch sử train/validation loss và metric theo epoch."),
            ("outputs/<experiment>/training_curves.png", "Biểu đồ train phục vụ báo cáo."),
            ("outputs/<experiment>/test_results.json", "Metric test chi tiết để đưa vào phần kết quả thí nghiệm."),
            ("outputs/<experiment>/test_results.png", "Hình phân bố distance, ROC curve và metric identification."),
            ("outputs/<experiment>/test_embeddings.pt", "Embedding, label, image_paths, gallery/probe indices và EER threshold cho demo sau này."),
        ],
        [6.3, 9.7],
    )

    doc.add_heading("8. Vai trò từng file mã nguồn", level=1)
    add_table(
        doc,
        ["File", "Vai trò"],
        [
            ("src/config.py", "Khai báo tham số dòng lệnh và kiểm tra cấu hình hợp lệ."),
            ("src/data.py", "Đọc VGGFace2, chia identity, tạo transform, dataset, PKBatchSampler và DataLoader."),
            ("src/model.py", "Định nghĩa PatchEmbedding, LayerScale, DropPath, TransformerEncoderBlock và FaceVisionTransformer."),
            ("src/train.py", "Huấn luyện bằng Semi-Hard Triplet Loss, lưu history và checkpoint."),
            ("src/train_infonce.py", "Huấn luyện bằng InfoNCE/Supervised Contrastive Loss."),
            ("src/infonce.py", "Cài đặt loss tương phản có temperature."),
            ("src/metrics.py", "Tính ROC-AUC, EER, TAR@FAR, Recall@K và mAP."),
            ("src/test.py", "Nạp checkpoint, trích xuất embedding, tạo pair/gallery/probe và lưu kết quả test."),
            ("src/visualization.py", "Vẽ biểu đồ train/test và lưu JSON kết quả."),
        ],
        [4.5, 11.5],
    )

    doc.add_heading("9. Kết quả và phần cần bổ sung", level=1)
    doc.add_paragraph(
        "Trong repo hiện tại chưa có thư mục outputs/checkpoints chứa kết quả training hoàn chỉnh, vì vậy báo cáo chưa thể điền số liệu VGGFace2 thực nghiệm. "
        "Khi chạy xong train.py hoặc train_infonce.py và test.py, cần bổ sung bảng kết quả từ test_results.json."
    )
    add_table(
        doc,
        ["Experiment", "Best epoch", "Val loss", "ROC-AUC", "EER", "Recall@1", "Recall@5", "mAP"],
        [
            ("sic_facevit_vggface2_semi_hard", "Chưa chạy", "Chưa có", "Chưa có", "Chưa có", "Chưa có", "Chưa có", "Chưa có"),
            ("supcon/infonce", "Chưa chạy", "Chưa có", "Chưa có", "Chưa có", "Chưa có", "Chưa có", "Chưa có"),
        ],
        [3.8, 1.8, 1.8, 1.8, 1.6, 1.8, 1.8, 1.6],
    )

    doc.add_heading("10. Kết luận", level=1)
    doc.add_paragraph(
        "Mô hình hiện tại là FaceViT tự xây dựng cho bài toán face embedding, không phải classifier closed-set. "
        "Thiết kế phù hợp với hướng open-set hơn vì người dùng ở giai đoạn test hoặc triển khai có thể là identity chưa xuất hiện trong train. "
        "Bước tiếp theo nên là chạy đầy đủ hai nhánh Triplet và InfoNCE trên cùng split, so sánh bằng ROC-AUC, EER, TAR@FAR, Recall@K và mAP, sau đó mới chọn checkpoint để tích hợp demo."
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
